import os
import fitz  # PyMuPDF
import docx

class DocumentExporter:
    @staticmethod
    def export(elements, output_path):
        """Elementleri alır ve istenen formatta yeni bir dosya olarak kaydeder."""
        ext = output_path.lower().split('.')[-1]

        try:
            if ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    for el in elements:
                        if el['type'] == 'text':
                            f.write(el['content'] + "\n\n")
            
            elif ext == 'docx':
                doc = docx.Document()
                for el in elements:
                    if el['type'] == 'text':
                        doc.add_paragraph(el['content'])
                    elif el['type'] == 'image':
                        # Görselleri geçici diske yazıp Word'e gömüyoruz
                        temp_img = f"vizia_temp_img.{el['ext']}"
                        with open(temp_img, "wb") as f:
                            f.write(el['data'])
                        doc.add_picture(temp_img)
                        os.remove(temp_img) # İzi yok et
                doc.save(output_path)
            
            elif ext == 'pdf':
                # Sınırsız sayfa destekleyen PDF oluşturucu
                doc = fitz.open()
                page = doc.new_page()
                y_pos = 50
                for el in elements:
                    if el['type'] == 'text':
                        rect = fitz.Rect(50, y_pos, 550, y_pos + 500)
                        rc = page.insert_textbox(rect, el['content'], fontsize=11, fontname="helv")
                        if rc < 0: # Eğer sayfa dolduysa yeni sayfa aç
                            page = doc.new_page()
                            y_pos = 50
                            page.insert_textbox(fitz.Rect(50, y_pos, 550, y_pos + 500), el['content'], fontsize=11)
                        y_pos += 40
                doc.save(output_path)
                
            return True
        except Exception as e:
            print(f"Yazma/Kaydetme Hatası: {e}")
            return False