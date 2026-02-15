import os
import fitz  # PyMuPDF
import docx

class DocumentEngine:
    @staticmethod
    def extract(file_path, include_images=False):
        """Dosyayı okur ve metin/görsel parçalarını (element listesi) döndürür."""
        ext = file_path.lower().split('.')[-1]
        elements = []

        try:
            if ext == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    elements.append({"type": "text", "content": f.read()})
            
            elif ext == 'docx':
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        elements.append({"type": "text", "content": para.text})
            
            elif ext == 'pdf':
                doc = fitz.open(file_path)
                for page in doc:
                    text = page.get_text()
                    if text.strip():
                        # Kırık satırları birleştirip anlam bütünlüğünü koruyoruz
                        clean_text = text.replace('-\n', '').replace('\n', ' ')
                        elements.append({"type": "text", "content": clean_text})
                    
                    if include_images:
                        for img in page.get_images():
                            base_image = doc.extract_image(img[0])
                            elements.append({"type": "image", "data": base_image["image"], "ext": base_image["ext"]})
        except Exception as e:
            print(f"Okuma Hatası: {e}")
            
        return elements

    @staticmethod
    def export(elements, output_path):
        """Elementleri alır ve istenen formatta yeni bir dosya olarak kaydeder."""
        ext = output_path.lower().split('.')[-1]

        try:
            if ext == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    for el in elements:
                        if el['type'] == 'text':
                            f.write(el['content'] + "\n\n")
            
            elif ext == 'docx':
                doc = docx.Document()
                for el in elements:
                    if el['type'] == 'text':
                        doc.add_paragraph(el['content'])
                    elif el['type'] == 'image':
                        # Görselleri geçici diske yazıp word'e gömüyoruz
                        temp_img = f"temp_img.{el['ext']}"
                        with open(temp_img, "wb") as f:
                            f.write(el['data'])
                        doc.add_picture(temp_img)
                        os.remove(temp_img)
                doc.save(output_path)
            
            elif ext == 'pdf':
                # Basit PDF oluşturucu. (Gelişmiş düzen için docx önerilir)
                doc = fitz.open()
                page = doc.new_page()
                y_pos = 50
                for el in elements:
                    if el['type'] == 'text':
                        rect = fitz.Rect(50, y_pos, 550, y_pos + 500)
                        rc = page.insert_textbox(rect, el['content'], fontsize=11, fontname="helv")
                        if rc < 0: # Sayfa dolduysa
                            page = doc.new_page()
                            y_pos = 50
                            page.insert_textbox(fitz.Rect(50, y_pos, 550, y_pos + 500), el['content'], fontsize=11)
                        y_pos += 40
                doc.save(output_path)
                
            return True
        except Exception as e:
            print(f"Yazma Hatası: {e}")
            return False